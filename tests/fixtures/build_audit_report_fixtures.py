"""Build deterministic, entirely synthetic report fixtures for extraction tests."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import zipfile

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas


ROOT = Path(__file__).resolve().parent / "generated"
STAMP = "SYNTHETIC TEST DATA"
FINDING = "Synthetic audit finding: approval was bypassed."
VERTICAL_BODY_SENTINEL = "VERTICAL-SYNTHETIC-FULL-BODY-SENTINEL-ONLY"
TEXT_PAGES = (
    "Synthetic audit report page one documents a simulated approval control gap. "
    "The example is fabricated for extraction testing and contains no real evidence.",
    "Synthetic audit report page two records a made-up remediation discussion. "
    "All names, events, conclusions, and amounts are fictional test material only.",
)


def _font(size: int) -> ImageFont.ImageFont:
    return ImageFont.truetype("arial.ttf", size)


def _finding_image(path: Path) -> None:
    image = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 55), STAMP, fill="black", font=_font(40))
    draw.text((70, 180), FINDING, fill="black", font=_font(48))
    draw.multiline_text(
        (70, 290),
        "This is synthetic evidence for automated extraction tests.\n"
        "No real audit report, person, company, or operational event is represented.",
        fill="black",
        font=_font(34),
        spacing=20,
    )
    image.save(path)


def _text_pdf(path: Path) -> None:
    canvas = Canvas(str(path), pagesize=letter, invariant=1)
    for text in TEXT_PAGES:
        canvas.setFont("Helvetica-Bold", 12)
        canvas.drawString(72, 740, STAMP)
        canvas.setFont("Helvetica", 12)
        y = 690
        for line in _wrapped_lines(canvas, text, letter[0] - 144):
            canvas.drawString(72, y, line)
            y -= 22
        canvas.showPage()
    canvas.save()


def _wrapped_lines(canvas: Canvas, text: str, max_width: float) -> list[str]:
    lines: list[str] = []
    current = ""
    for word in text.split():
        candidate = word if not current else f"{current} {word}"
        if current and canvas.stringWidth(candidate, "Helvetica", 12) > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def _scan_pdf(path: Path, raster: Path) -> None:
    canvas = Canvas(str(path), pagesize=letter, invariant=1)
    canvas.drawImage(str(raster), 42, 220, width=528, height=251)
    canvas.save()


def _vertical_raster(path: Path) -> None:
    """Large, high-contrast raster text for the real local RapidOCR path."""
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((90, 80), STAMP, fill="black", font=_font(56))
    draw.text((90, 250), "LOCATOR GAMMA", fill="black", font=_font(72))
    draw.text((90, 410), "Fictional scanned control exception.", fill="black", font=_font(48))
    draw.text((90, 540), "No real entity person or event.", fill="black", font=_font(42))
    image.save(path, format="PNG", optimize=False)


def _vertical_pdf(path: Path, raster: Path) -> None:
    """A fixed three-page PDF: two text pages and one raster-only scan."""
    canvas = Canvas(str(path), pagesize=letter, invariant=1)
    pages = (
        ("LOCATOR ALPHA", "Fictional approval exception for local test only."),
        ("LOCATOR BETA", "Fictional review exception for local test only."),
    )
    for locator, sentence in pages:
        canvas.setFont("Helvetica-Bold", 18)
        canvas.drawString(72, 735, STAMP)
        canvas.setFont("Helvetica-Bold", 16)
        canvas.drawString(72, 680, locator)
        canvas.setFont("Helvetica", 12)
        canvas.drawString(72, 640, sentence)
        canvas.drawString(72, 600, VERTICAL_BODY_SENTINEL)
        canvas.showPage()
    canvas.drawImage(str(raster), 42, 180, width=528, height=264)
    canvas.showPage()
    canvas.save()


def _stable_docx(path: Path) -> None:
    with zipfile.ZipFile(path) as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name, data in sorted(entries):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.create_system = 0
            info.external_attr = 0
            info.compress_type = zipfile.ZIP_DEFLATED
            target.writestr(info, data, compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def _document() -> Document:
    document = Document()
    fixed_time = datetime(1980, 1, 1, tzinfo=timezone.utc)
    document.core_properties.subject = "虚构测试资料 / SYNTHETIC TEST DATA"
    document.core_properties.comments = "Synthetic fixture; no real audit report."
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.core_properties.last_modified_by = "Synthetic Fixture Builder"
    document.core_properties.revision = 1
    return document


def _docx(path: Path, image: Path, mixed: bool) -> None:
    document = _document()
    document.add_heading("虚构测试资料 / SYNTHETIC TEST DATA", level=1)
    if mixed:
        document.add_paragraph("Synthetic opening paragraph before the embedded finding.")
        document.add_picture(str(image), width=Inches(5.8))
        document.add_paragraph("Synthetic closing paragraph after the embedded finding.")
    else:
        document.add_paragraph("Synthetic first paragraph keeps document order visible.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Synthetic control"
        table.cell(0, 1).text = "Synthetic status"
        table.cell(1, 0).text = "Approval review"
        table.cell(1, 1).text = "Bypassed in fictional sample"
        document.add_paragraph("Synthetic final paragraph follows the table in document order.")
    document.save(path)
    _stable_docx(path)


def _inline_docx(path: Path, image: Path) -> None:
    document = _document()
    document.add_heading("虚构测试资料 / SYNTHETIC TEST DATA", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("INLINE BEFORE")
    paragraph.add_run().add_picture(str(image), width=Inches(4.5))
    paragraph.add_run("INLINE AFTER")
    document.save(path)
    _stable_docx(path)


def _table_image_docx(path: Path, image: Path) -> None:
    document = _document()
    document.add_heading("虚构测试资料 / SYNTHETIC TEST DATA", level=1)
    table = document.add_table(rows=1, cols=2)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.add_run("CELL BEFORE")
    paragraph.add_run().add_picture(str(image), width=Inches(2.3))
    paragraph.add_run("CELL AFTER")
    table.cell(0, 1).text = "SECOND CELL"
    document.save(path)
    _stable_docx(path)


def build() -> Path:
    ROOT.mkdir(parents=True, exist_ok=True)
    raster = ROOT / "synthetic_finding.png"
    _finding_image(raster)
    _text_pdf(ROOT / "text_report.pdf")
    _scan_pdf(ROOT / "scan_report.pdf", raster)
    vertical_raster = ROOT / "vertical_slice_scan.png"
    _vertical_raster(vertical_raster)
    _vertical_pdf(ROOT / "vertical_slice_report.pdf", vertical_raster)
    _docx(ROOT / "report.docx", raster, mixed=False)
    _docx(ROOT / "mixed_report.docx", raster, mixed=True)
    _inline_docx(ROOT / "inline_report.docx", raster)
    _table_image_docx(ROOT / "table_image_report.docx", raster)
    return ROOT


if __name__ == "__main__":
    build()
