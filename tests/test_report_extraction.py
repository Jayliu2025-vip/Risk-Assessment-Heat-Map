from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from pathlib import Path
import shutil
import tempfile
import unittest
from unittest.mock import patch
import zipfile

import desktop.extraction as extraction
from desktop.extraction import ExtractionError, extract_docx_source_text, extract_report, text_is_usable
from desktop.ocr import OcrError, RapidOcrEngine
from reportlab.pdfgen.canvas import Canvas
from docx import Document
from tests.fixtures.build_audit_report_fixtures import TEXT_PAGES, _wrapped_lines, build


FIXTURES = Path(__file__).parent / "fixtures" / "generated"


class FakeOcr:
    def __init__(self, text: str = "Synthetic audit finding: approval was bypassed."):
        self.text = text
        self.calls: list[Path] = []

    def read(self, image_path: Path) -> str:
        self.calls.append(Path(image_path))
        return self.text


class ReportExtractionTests(unittest.TestCase):
    def setUp(self):
        self.temp = Path(tempfile.mkdtemp(prefix="report-extraction-"))

    def tearDown(self):
        shutil.rmtree(self.temp, ignore_errors=True)

    def test_text_is_usable_rejects_blank_short_and_replacement_text(self):
        self.assertFalse(text_is_usable(" \n\t "))
        self.assertFalse(text_is_usable("short text"))
        self.assertFalse(text_is_usable("a" * 41 + "\ufffd" * 10))
        self.assertTrue(text_is_usable("a" * 40))

    def test_text_pdf_has_two_text_blocks_without_ocr(self):
        ocr = FakeOcr()
        result = extract_report(FIXTURES / "text_report.pdf", self.temp, ocr)
        self.assertEqual(result.method, "text")
        self.assertEqual([block.locator for block in result.blocks], ["第 1 页", "第 2 页"])
        self.assertTrue(all(block.method == "text" for block in result.blocks))
        self.assertEqual(ocr.calls, [])

    def test_scanned_pdf_routes_page_through_ocr(self):
        ocr = FakeOcr()
        result = extract_report(FIXTURES / "scan_report.pdf", self.temp, ocr)
        self.assertEqual(result.method, "ocr")
        self.assertEqual(len(ocr.calls), 1)
        self.assertEqual(result.blocks[0].locator, "第 1 页")
        self.assertEqual(result.blocks[0].method, "ocr")

    def test_failed_ocr_routes_to_vision_and_keeps_contained_render(self):
        result = extract_report(FIXTURES / "scan_report.pdf", self.temp, FakeOcr("too short"))
        block = result.blocks[0]
        self.assertEqual(result.method, "vision_required")
        self.assertEqual(block.method, "vision_required")
        self.assertTrue(block.needs_review)
        self.assertIsNotNone(block.image_path)
        image = Path(block.image_path)
        self.assertTrue(image.is_file())
        self.assertIn(self.temp.resolve(), image.resolve().parents)

    def test_docx_preserves_paragraph_table_order_and_serializes_cells(self):
        result = extract_report(FIXTURES / "report.docx", self.temp, FakeOcr())
        self.assertEqual(result.method, "text")
        self.assertEqual([block.locator for block in result.blocks], ["Word 段落 1", "Word 段落 2", "Word 表格 1", "Word 段落 3"])
        self.assertIn("Synthetic control\tSynthetic status", result.blocks[2].text)
        self.assertIn("Approval review\tBypassed", result.blocks[2].text)

    def test_mixed_docx_routes_embedded_image_in_document_order(self):
        ocr = FakeOcr()
        result = extract_report(FIXTURES / "mixed_report.docx", self.temp, ocr)
        self.assertEqual(result.method, "mixed")
        self.assertEqual([block.method for block in result.blocks], ["text", "text", "ocr", "text"])
        self.assertEqual(result.blocks[2].locator, "Word 段落 3")
        self.assertEqual(len(ocr.calls), 1)

    def test_unsupported_and_corrupt_inputs_have_stable_safe_errors(self):
        legacy = self.temp / "old.doc"
        legacy.write_bytes(b"synthetic")
        unsupported = self.temp / "input.txt"
        unsupported.write_text("synthetic", encoding="utf-8")
        corrupt_pdf = self.temp / "bad.pdf"
        corrupt_pdf.write_bytes(b"not a pdf")
        corrupt_docx = self.temp / "bad.docx"
        corrupt_docx.write_bytes(b"not a docx")
        cases = ((legacy, "DOC_LEGACY_UNSUPPORTED"), (unsupported, "FILE_TYPE_UNSUPPORTED"), (corrupt_pdf, "PDF_READ_FAILED"), (corrupt_docx, "DOCX_READ_FAILED"))
        for path, code in cases:
            with self.subTest(path=path.name):
                with self.assertRaises(ExtractionError) as caught:
                    extract_report(path, self.temp, FakeOcr())
                self.assertEqual(caught.exception.code, code)
                self.assertNotIn(str(path.parent), caught.exception.message)

    def test_temp_root_must_contain_generated_outputs(self):
        outside = self.temp / "missing" / "nested"
        with self.assertRaises(ExtractionError) as caught:
            extract_report(FIXTURES / "scan_report.pdf", outside, FakeOcr())
        self.assertEqual(caught.exception.code, "TEMP_DIR_INVALID")

    def test_rapid_ocr_engine_joins_injected_result_texts(self):
        @dataclass
        class Result:
            txts: list[str]

        engine = RapidOcrEngine(engine=lambda _: Result(["alpha", "beta"]))
        self.assertEqual(engine.read(Path("synthetic.png")), "alpha\nbeta")

    def test_real_local_ocr_smoke_reads_synthetic_scan(self):
        result = extract_report(FIXTURES / "scan_report.pdf", self.temp, RapidOcrEngine())
        normalized = " ".join(result.blocks[0].text.lower().split())
        self.assertIn("synthetic", normalized)
        self.assertIn("approval", normalized)

    def test_source_size_limit_is_checked_before_opening(self):
        with patch.object(extraction, "MAX_SOURCE_BYTES", 1):
            with self.assertRaises(ExtractionError) as caught:
                extract_report(FIXTURES / "text_report.pdf", self.temp, FakeOcr())
        self.assertEqual(caught.exception.code, "SOURCE_TOO_LARGE")
        self.assertNotIn(str(FIXTURES), caught.exception.message)

    def test_pdf_page_and_render_limits_are_rejected_before_ocr(self):
        with patch.object(extraction, "MAX_PDF_PAGES", 0):
            with self.assertRaises(ExtractionError) as caught:
                extract_report(FIXTURES / "text_report.pdf", self.temp, FakeOcr())
        self.assertEqual(caught.exception.code, "PDF_PAGE_LIMIT")
        with patch.object(extraction, "MAX_RENDER_PIXELS", 1):
            with self.assertRaises(ExtractionError) as caught:
                extract_report(FIXTURES / "scan_report.pdf", self.temp, FakeOcr())
        self.assertEqual(caught.exception.code, "PDF_RENDER_LIMIT")

    def test_docx_zip_preflight_rejects_entry_count_size_and_ratio(self):
        cases = (("MAX_DOCX_ENTRIES", 1, "DOCX_ENTRY_LIMIT"), ("MAX_DOCX_UNCOMPRESSED_BYTES", 1, "DOCX_SIZE_LIMIT"), ("MAX_DOCX_COMPRESSION_RATIO", 1, "DOCX_COMPRESSION_RATIO"))
        for constant, value, code in cases:
            with self.subTest(constant=constant), patch.object(extraction, constant, value):
                with self.assertRaises(ExtractionError) as caught:
                    extract_report(FIXTURES / "report.docx", self.temp, FakeOcr())
                self.assertEqual(caught.exception.code, code)

    def test_docx_zip_preflight_rejects_duplicate_and_traversal_entries(self):
        duplicate = self.temp / "duplicate.docx"
        traversal = self.temp / "traversal.docx"
        with zipfile.ZipFile(FIXTURES / "report.docx") as source:
            entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
        for target, extra in ((duplicate, ("word/./document.xml", b"synthetic")), (traversal, ("../escape", b"synthetic"))):
            with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as archive:
                for name, data in entries:
                    archive.writestr(name, data)
                archive.writestr(*extra)
        for path, code in ((duplicate, "DOCX_DUPLICATE_ENTRY"), (traversal, "DOCX_ENTRY_PATH_INVALID")):
            with self.subTest(path=path.name):
                with self.assertRaises(ExtractionError) as caught:
                    extract_report(path, self.temp, FakeOcr())
                self.assertEqual(caught.exception.code, code)

    def test_docx_image_byte_and_pixel_limits_are_checked_before_ocr(self):
        for constant, code in (("MAX_IMAGE_BYTES", "IMAGE_BYTES_LIMIT"), ("MAX_IMAGE_PIXELS", "IMAGE_PIXELS_LIMIT")):
            with self.subTest(constant=constant), patch.object(extraction, constant, 1):
                with self.assertRaises(ExtractionError) as caught:
                    extract_report(FIXTURES / "mixed_report.docx", self.temp, FakeOcr())
                self.assertEqual(caught.exception.code, code)

    def test_inline_docx_image_preserves_before_image_after_order(self):
        result = extract_report(FIXTURES / "inline_report.docx", self.temp, FakeOcr())
        texts = [block.text for block in result.blocks if block.locator == "Word 段落 2"]
        self.assertEqual(texts, ["INLINE BEFORE", "Synthetic audit finding: approval was bypassed.", "INLINE AFTER"])
        self.assertEqual([block.method for block in result.blocks if block.locator == "Word 段落 2"], ["text", "ocr", "text"])

    def test_table_cell_image_preserves_cell_text_and_image_order(self):
        result = extract_report(FIXTURES / "table_image_report.docx", self.temp, FakeOcr())
        table_blocks = [block for block in result.blocks if block.locator == "Word 表格 1"]
        self.assertEqual([block.method for block in table_blocks], ["text", "ocr", "text"])
        self.assertEqual([block.text for block in table_blocks], ["CELL BEFORE", "Synthetic audit finding: approval was bypassed.", "CELL AFTER\tSECOND CELL"])

    def test_docx_source_preview_reuses_ocr_blocks_for_image_only_and_inline_locations(self):
        image_only = extract_docx_source_text(
            FIXTURES / "mixed_report.docx", "Word 段落 3", self.temp, FakeOcr("IMAGE ONLY OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE"))
        inline = extract_docx_source_text(
            FIXTURES / "inline_report.docx", "Word 段落 2", self.temp, FakeOcr("INLINE OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE"))
        table = extract_docx_source_text(
            FIXTURES / "table_image_report.docx", "Word 表格 1", self.temp, FakeOcr("TABLE OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE"))
        self.assertEqual(image_only, "IMAGE ONLY OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE")
        self.assertEqual(inline.splitlines(), ["INLINE BEFORE", "INLINE OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE", "INLINE AFTER"])
        self.assertEqual(table.splitlines(), ["CELL BEFORE", "TABLE OCR EVIDENCE FROM SYNTHETIC AUDIT SOURCE", "CELL AFTER\tSECOND CELL"])

    def test_ocr_import_init_and_inference_failures_are_sanitized(self):
        for factory in (lambda: (_ for _ in ()).throw(ImportError("secret path")), lambda: (_ for _ in ()).throw(RuntimeError("secret text"))):
            with self.subTest(factory=factory):
                with self.assertRaises(OcrError) as caught:
                    RapidOcrEngine(factory=factory).read(Path("private.png"))
                self.assertEqual(caught.exception.code, "OCR_UNAVAILABLE")
                self.assertNotIn("secret", caught.exception.message)
        with self.assertRaises(OcrError) as caught:
            RapidOcrEngine(engine=lambda _: (_ for _ in ()).throw(RuntimeError("secret text"))).read(Path("private.png"))
        self.assertEqual(caught.exception.code, "OCR_FAILED")
        self.assertNotIn("secret", caught.exception.message)

    def test_ocr_failure_routes_scan_to_vision_without_leaking_error_text(self):
        ocr = RapidOcrEngine(engine=lambda _: (_ for _ in ()).throw(RuntimeError("private evidence")))
        result = extract_report(FIXTURES / "scan_report.pdf", self.temp, ocr)
        self.assertEqual(result.blocks[0].method, "vision_required")
        self.assertTrue(result.blocks[0].needs_review)
        self.assertNotIn("private", result.blocks[0].text)
        self.assertTrue(Path(result.blocks[0].image_path).is_file())

    def test_docx_builder_is_deterministic_for_both_docx_fixtures(self):
        build()
        first = {name: sha256((FIXTURES / name).read_bytes()).hexdigest() for name in ("report.docx", "mixed_report.docx")}
        build()
        second = {name: sha256((FIXTURES / name).read_bytes()).hexdigest() for name in first}
        self.assertEqual(first, second)

    def test_pdf_fixture_word_wrap_keeps_every_line_inside_text_area(self):
        canvas = Canvas(BytesIO())
        canvas.setFont("Helvetica", 12)
        for page in TEXT_PAGES:
            lines = _wrapped_lines(canvas, page, 468)
            self.assertEqual(" ".join(lines), page)
            self.assertTrue(all(canvas.stringWidth(line, "Helvetica", 12) <= 468 for line in lines))

    def test_table_cell_paragraphs_keep_a_single_newline_boundary(self):
        source = self.temp / "synthetic_multi_paragraph_table.docx"
        document = Document()
        document.add_heading("虚构测试资料 / SYNTHETIC TEST DATA", level=1)
        cell = document.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "FIRST"
        cell.add_paragraph("SECOND")
        document.save(source)
        result = extract_report(source, self.temp, FakeOcr())
        text = "".join(block.text for block in result.blocks if block.locator == "Word 表格 1")
        self.assertIn("FIRST\nSECOND", text)
        self.assertNotIn("FIRSTSECOND", text)


if __name__ == "__main__":
    unittest.main()
