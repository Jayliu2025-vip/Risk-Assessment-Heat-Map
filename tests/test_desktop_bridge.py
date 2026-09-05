"""Synthetic contract tests for the pywebview desktop bridge."""

from __future__ import annotations

from dataclasses import asdict
import hashlib
import inspect
from pathlib import Path
import shutil
import tempfile
import threading
import unittest
from unittest.mock import patch

from docx import Document

from desktop.models import AnalysisTask, FindingDraft, ModelProfile
from desktop.storage import DesktopStore
from tools.common import DIMS


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def finding(task_id: str, finding_id: str = "F-1", **changes):
    value = FindingDraft(
        task_id=task_id, finding_id=finding_id, title="合成发现", fact_summary="合成事实",
        source_page="第 1 页", source_excerpt="合成摘录", matched_risk_id="R001",
        domain="采购与外包", likelihood=3, impact_scores={dim: 2 for dim in DIMS},
        rationale="合成依据", needs_review=True,
    )
    for key, item in changes.items():
        setattr(value, key, item)
    return value


class Window:
    def __init__(self, selected: list[str] | None = None):
        self.selected = selected or []
        self.calls = []

    def create_file_dialog(self, **kwargs):
        self.calls.append(kwargs)
        return self.selected


class CredentialMemory:
    def __init__(self): self.values = {}
    def set_password(self, service, name, secret): self.values[(service, name)] = secret
    def get_password(self, service, name): return self.values.get((service, name))
    def set_api_key(self, name, secret): self.values[name] = secret
    def get_api_key(self, name): return self.values.get(name)
    def delete_api_key(self, name): self.values.pop(name, None)


class Pipeline:
    def __init__(self, store): self.store, self.calls, self.cleaned = store, [], []
    def start(self, source_path, model_profile, risk_catalog):
        self.calls.append((Path(source_path), model_profile, list(risk_catalog)))
        task = AnalysisTask("T-1", Path(source_path).name, digest(Path(source_path)), "2026-09-03T00:00:00Z", "待复核", model_profile, "text")
        self.store.save_task(task)
        return task
    def events(self, task_id): return [{"status": "待复核", "code": "DONE", "message": "完成", "timestamp": "x"}]
    def review_findings(self, task_id, edits):
        checked = [item if isinstance(item, FindingDraft) else FindingDraft(**item) for item in edits]
        return self.store.save_findings(checked)
    def cleanup_task(self, task_id): self.cleaned.append(task_id)


class Writer:
    def __init__(self): self.preview_calls, self.commit_calls = [], []
    def preview_changes(self, source, decisions, findings):
        self.preview_calls.append((Path(source), decisions, findings))
        return {"commit_token": "token-1", "new_risks": [], "updated_risks": [],
                "new_controls": [], "excluded_count": 0, "warnings": []}
    def write_versioned_workbook(self, source, decisions, findings, *, expected_commit_token):
        self.commit_calls.append((Path(source), decisions, findings, expected_commit_token))
        return type("Result", (), {"workbook_path": Path("C:/synthetic/out.xlsx"), "export_dir": Path("C:/synthetic/export"), "periods": ["2026H1"], "assessed_risks": [{"risk_id": "R001"}]})()


class Client:
    def __init__(self, profile, key): self.profile, self.key = profile, key
    def test_connection(self): return True
    def close(self): pass


class DesktopBridgeTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.store = DesktopStore(self.root / "state.db")
        self.pipeline = Pipeline(self.store)
        self.writer = Writer()
        self.window = Window()
        self.report = self.root / "synthetic.pdf"
        self.report.write_bytes(b"synthetic-pdf")
        self.workbook = self.root / "synthetic.xlsx"
        shutil.copy2(Path(__file__).resolve().parents[1] / "audit_risk_register.xlsx", self.workbook)
        from desktop.bridge import DesktopBridge
        self.bridge = DesktopBridge(
            store=self.store, pipeline=self.pipeline, credential_store=CredentialMemory(),
            model_client_factory=Client, workbook_writer=self.writer,
            risk_catalog=[{"risk_id": "R001", "name": "合成风险"}],
            pdf_preview_renderer=lambda path, page: b"png-bytes",
        )
        self.bridge.attach_window(self.window)

    def tearDown(self):
        self.store.close()
        self.temp.cleanup()

    def profile(self):
        return {"name": "synthetic", "base_url": "https://model.example.test", "model": "test", "supports_vision": False, "api_key": "sk-synthetic-secret"}

    def select(self, path, purpose="report"):
        self.window.selected = [str(path)]
        return self.bridge.choose_report(purpose)

    def seed_task(self, task_id="T-1", source=None):
        source = source or self.report
        self.store.save_task(AnalysisTask(task_id, source.name, digest(source), "2026-09-03T00:00:00Z", "待复核", "synthetic", "text"))
        self.store.save_findings([finding(task_id)])

    def create_decision(self, finding_id="F-1", period="2026H1"):
        return {"action": "create", "finding_ids": [finding_id], "risk_id": "",
                "name": "合成风险", "domain": "采购与外包", "description": "合成事实",
                "owner_dept": "审计部", "period": period, "likelihood": 3,
                "impact_scores": {dim: 2 for dim in DIMS}, "rationale": "合成依据",
                "remediation_status": "未整改", "controls": []}

    def bind_workbook(self, task_id="T-1", period="2026H1", path=None):
        path = path or self.workbook
        token = self.select(path, "workbook")["selection_token"]
        from desktop.workbook_writer import load_risk_catalog
        self.bridge._task_workbooks[task_id] = (path.resolve(), token, digest(path), period, load_risk_catalog(path))
        return token

    def test_exact_js_public_allowlist(self):
        actual = {name for name, value in inspect.getmembers(type(self.bridge), inspect.isfunction) if not name.startswith("_")}
        self.assertEqual(actual, {
            "get_bootstrap", "choose_report", "choose_catalog_root", "configure_workspace",
            "list_catalog_reports", "save_report_to_catalog", "trash_catalog_report",
            "clear_catalog_reports", "create_catalog_batch", "get_source_preview",
            "save_model_profile", "test_model_profile", "start_analysis", "get_task",
            "get_findings", "save_finding", "merge_findings", "split_finding",
            "preview_commit", "commit_to_workbook", "cleanup_task",
        })

    def test_choose_report_returns_token_only_and_rejects_wrong_purpose(self):
        chosen = self.select(self.report)
        self.assertEqual(chosen["ok"], True)
        self.assertEqual(chosen["basename"], "synthetic.pdf")
        self.assertNotIn(str(self.report), str(chosen))
        wrong = self.bridge.start_analysis(chosen["selection_token"], "missing", "2026H1", "synthetic")
        self.assertFalse(wrong["ok"])
        self.assertEqual(wrong["code"], "SELECTION_NOT_FOUND")
        workbook = self.select(self.workbook, "workbook")
        self.assertEqual(self.bridge.start_analysis(workbook["selection_token"], workbook["selection_token"], "2026H1", "synthetic")["code"], "SELECTION_PURPOSE_INVALID")
        self.assertEqual(self.bridge.start_analysis("stale", workbook["selection_token"], "2026H1", "synthetic")["code"], "SELECTION_NOT_FOUND")

    def test_profile_never_echoes_secret_and_connection_returns_hostname(self):
        saved = self.bridge.save_model_profile(self.profile())
        self.assertTrue(saved["ok"])
        self.assertNotIn("api_key", saved)
        self.assertNotIn("sk-synthetic-secret", str(saved))
        tested = self.bridge.test_model_profile("synthetic")
        self.assertEqual(tested, {"ok": True, "hostname": "model.example.test"})
        self.assertNotIn("sk-synthetic-secret", str(self.bridge.get_bootstrap()))

    def test_profile_key_pair_rolls_back_if_profile_persistence_fails(self):
        self.bridge.save_model_profile(self.profile())
        old = self.store.list_model_profiles()[0]
        self.store.save_model_profile = lambda profile: (_ for _ in ()).throw(RuntimeError("C:\\secret\\sk-new"))
        changed = {**self.profile(), "base_url": "https://new.example.test", "api_key": "sk-new"}
        result = self.bridge.save_model_profile(changed)
        self.assertFalse(result["ok"])
        self.assertEqual(self.bridge._credential_store.get_api_key("synthetic"), "sk-synthetic-secret")
        self.assertEqual(old.base_url, "https://model.example.test")
        self.assertNotIn("sk-new", str(result))

    def test_start_task_findings_are_serializable_and_never_source_path(self):
        self.bridge.save_model_profile(self.profile())
        self.bridge.test_model_profile("synthetic")
        selected = self.select(self.report)
        workbook = self.select(self.workbook, "workbook")
        started = self.bridge.start_analysis(selected["selection_token"], workbook["selection_token"], "2026H1", "synthetic")
        self.assertTrue(started["ok"])
        self.assertNotIn(str(self.report), str(started))
        self.store.save_findings([finding("T-1")])
        self.assertEqual(self.bridge.get_task("T-1")["task"]["task_id"], "T-1")
        self.assertEqual(self.bridge.get_findings("T-1")["findings"][0]["finding_id"], "F-1")

    def test_error_sanitizer_hides_secret_path_and_report_body(self):
        self.bridge._pipeline = type("Bad", (), {"start": lambda *_: (_ for _ in ()).throw(RuntimeError("sk-secret C:\\secret\\report.pdf 完整报告正文"))})()
        result = self.bridge.start_analysis("missing", "missing", "2026H1", "synthetic")
        self.assertFalse(result["ok"])
        for forbidden in ("sk-secret", "C:\\secret", "完整报告正文"):
            self.assertNotIn(forbidden, str(result))

    def test_source_preview_hash_pdf_docx_and_reselect_contract(self):
        self.seed_task()
        selected = self.select(self.report)
        self.bridge._task_sources["T-1"] = (self.report, selected["selection_token"])
        pdf = self.bridge.get_source_preview("T-1", "F-1")
        self.assertEqual(pdf["kind"], "pdf")
        self.assertTrue(pdf["image_data_url"].startswith("data:image/png;base64,"))
        self.report.write_bytes(b"changed")
        self.assertEqual(self.bridge.get_source_preview("T-1", "F-1")["code"], "SOURCE_HASH_CHANGED")
        self.assertEqual(self.bridge.get_source_preview("other", "F-1")["code"], "SOURCE_RESELECT_REQUIRED")
        docx = self.root / "synthetic.docx"
        document = Document()
        document.add_paragraph("第一段虚构内容")
        document.add_paragraph("第二段来自哈希校验后的 Word 源文件，包含虚构单位甲和金额 12 万元。")
        document.save(docx)
        self.seed_task("T-docx", docx)
        self.store.save_findings([finding("T-docx", "F-docx", source_page="Word 段落 2", source_excerpt="模型伪造摘录，不得回显")])
        token = self.select(docx)["selection_token"]
        self.bridge._task_sources["T-docx"] = (docx, token)
        preview = self.bridge.get_source_preview("T-docx", "F-docx")
        self.assertEqual(preview, {"ok": True, "kind": "text", "source_page": "Word 段落 2",
                                   "source_excerpt": "第二段来自哈希校验后的 Word 源文件，包含虚构单位甲和金额 12 万元。"})
        self.assertNotIn("模型伪造摘录", str(preview))

    def test_docx_preview_parser_is_injectable_and_receives_private_snapshot(self):
        docx = self.root / "source.docx"
        document = Document(); document.add_paragraph("虚构正文"); document.save(docx)
        self.seed_task("T-docx-injected", docx)
        self.store.save_findings([finding("T-docx-injected", "F-docx", source_page="Word 段落 1")])
        seen = []
        self.bridge._docx_preview_extractor = lambda path, locator, temp_root: (seen.append((Path(path), locator, Path(temp_root))) or "重新提取文本")
        token = self.select(docx)["selection_token"]
        self.bridge._task_sources["T-docx-injected"] = (docx, token)
        preview = self.bridge.get_source_preview("T-docx-injected", "F-docx")
        self.assertEqual(preview["source_excerpt"], "重新提取文本")
        self.assertNotEqual(seen[0][0], docx)
        self.assertEqual(seen[0][1], "Word 段落 1")

    def test_model_test_is_bound_to_current_profile_and_credential(self):
        self.bridge.save_model_profile(self.profile())
        report = self.select(self.report)["selection_token"]
        workbook = self.select(self.workbook, "workbook")["selection_token"]
        self.assertEqual(self.bridge.start_analysis(report, workbook, "2026H1", "synthetic")["code"], "MODEL_PROFILE_TEST_REQUIRED")
        self.bridge.test_model_profile("synthetic")
        self.assertTrue(self.bridge.start_analysis(report, workbook, "2026H1", "synthetic")["ok"])
        changed = {**self.profile(), "model": "changed-model", "api_key": "sk-new"}
        self.bridge.save_model_profile(changed)
        self.assertEqual(self.bridge.start_analysis(report, workbook, "2026H1", "synthetic")["code"], "MODEL_PROFILE_TEST_REQUIRED")

    def test_analysis_catalog_and_later_commit_are_bound_to_selected_workbook_hash(self):
        self.bridge.save_model_profile(self.profile()); self.bridge.test_model_profile("synthetic")
        report = self.select(self.report)["selection_token"]
        workbook = self.select(self.workbook, "workbook")["selection_token"]
        started = self.bridge.start_analysis(report, workbook, "2026H1", "synthetic")
        self.assertTrue(started["ok"])
        model_catalog = self.pipeline.calls[-1][2]
        row = next(item for item in model_catalog if item["risk_id"] == "R001")
        self.assertEqual(set(row), {"risk_id", "name", "domain", "description"})
        local_row = next(item for item in started["risk_catalog"] if item["risk_id"] == "R001" and item["period"] == "2026H1")
        self.assertEqual(local_row["owner_dept"], "财务部")
        self.workbook.write_bytes(b"changed after analysis")
        blocked = self.bridge.preview_commit("T-1", workbook, "2026H1", [self.create_decision()], "preview", True)
        self.assertEqual(blocked["code"], "WORKBOOK_HASH_CHANGED")

    def test_pdf_preview_renders_private_snapshot_and_rejects_oversized_png(self):
        self.seed_task()
        selected = self.select(self.report)
        seen = []
        def renderer(path, page):
            seen.append(Path(path))
            self.report.write_bytes(b"replaced after snapshot")
            return b"small-png"
        self.bridge._pdf_preview_renderer = renderer
        self.bridge._task_sources["T-1"] = (self.report, selected["selection_token"])
        result = self.bridge.get_source_preview("T-1", "F-1")
        self.assertTrue(result["ok"])
        self.assertNotEqual(seen[0], self.report)
        import desktop.bridge as bridge_module
        previous = bridge_module.MAX_PREVIEW_PNG_BYTES
        bridge_module.MAX_PREVIEW_PNG_BYTES = 3
        try:
            self.report.write_bytes(b"synthetic-pdf")
            self.assertEqual(self.bridge.get_source_preview("T-1", "F-1")["code"], "PREVIEW_TOO_LARGE")
        finally:
            bridge_module.MAX_PREVIEW_PNG_BYTES = previous

    def test_preview_busy_is_safe_and_commit_token_is_consumed_once(self):
        self.seed_task()
        selected = self.select(self.report)
        self.assertTrue(self.bridge._preview_slots.acquire(blocking=False))
        try:
            self.assertEqual(self.bridge.get_source_preview("T-1", "F-1", selected["selection_token"])["code"], "PREVIEW_BUSY")
        finally:
            self.bridge._preview_slots.release()
        selected = {"selection_token": self.bind_workbook()}
        decision = [self.create_decision()]
        preview = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", decision, "preview", True)
        with patch("desktop.bridge.load_dataset", return_value=({}, [{"risk_id": "R001", "name": "合成风险", "domain": "采购与外包", "description": "合成事实", "owner_dept": "审计部", "period": "2026H1", "likelihood": 3, **{dim: 2 for dim in DIMS}, "rationale": "合成依据"}], [])):
            first = self.bridge.commit_to_workbook("T-1", selected["selection_token"], "2026H1", decision, preview["commit_token"])
        self.assertTrue(first["ok"])
        second = self.bridge.commit_to_workbook("T-1", selected["selection_token"], "2026H1", decision, preview["commit_token"])
        self.assertEqual(second["code"], "PREVIEW_REQUIRED")
        self.assertEqual(len(self.writer.commit_calls), 1)

    def test_source_preview_reattaches_a_matching_report_after_bridge_restart(self):
        self.seed_task()
        restarted = type(self.bridge)(
            store=self.store, pipeline=Pipeline(self.store), credential_store=CredentialMemory(),
            model_client_factory=Client, workbook_writer=self.writer,
            risk_catalog=[{"risk_id": "R001", "name": "合成风险"}],
            pdf_preview_renderer=lambda path, page: b"png-bytes",
        )
        restarted.attach_window(self.window)
        self.window.selected = [str(self.report)]
        token = restarted.choose_report("report")["selection_token"]
        preview = restarted.get_source_preview("T-1", "F-1", token)
        self.assertEqual(preview["kind"], "pdf")
        self.assertEqual(restarted._task_sources["T-1"][1], token)
        wrong = self.root / "other.pdf"; wrong.write_bytes(b"different synthetic report")
        mismatched = type(self.bridge)(
            store=self.store, pipeline=Pipeline(self.store), credential_store=CredentialMemory(),
            model_client_factory=Client, workbook_writer=self.writer,
            risk_catalog=[{"risk_id": "R001", "name": "合成风险"}],
            pdf_preview_renderer=lambda path, page: b"png-bytes",
        )
        mismatched.attach_window(self.window)
        self.window.selected = [str(wrong)]
        wrong_token = mismatched.choose_report("report")["selection_token"]
        mismatch = mismatched.get_source_preview("T-1", "F-1", wrong_token)
        self.assertEqual(mismatch["code"], "SOURCE_HASH_CHANGED")
        self.assertNotIn(str(wrong), str(mismatch))

    def test_preview_reattachment_is_serialized_with_cleanup(self):
        self.seed_task()
        restarted = type(self.bridge)(
            store=self.store, pipeline=Pipeline(self.store), credential_store=CredentialMemory(),
            model_client_factory=Client, workbook_writer=self.writer,
            risk_catalog=[{"risk_id": "R001", "name": "合成风险"}],
            pdf_preview_renderer=lambda path, page: b"png-bytes",
        )
        restarted.attach_window(self.window)
        self.window.selected = [str(self.report)]
        token = restarted.choose_report("report")["selection_token"]
        entered, release = threading.Event(), threading.Event()
        restarted._preview_before_attach = lambda: (entered.set(), release.wait(2))
        preview_result, cleanup_result = [], []
        preview_thread = threading.Thread(target=lambda: preview_result.append(restarted.get_source_preview("T-1", "F-1", token)))
        preview_thread.start()
        self.assertTrue(entered.wait(1))
        cleanup_thread = threading.Thread(target=lambda: cleanup_result.append(restarted.cleanup_task("T-1")))
        cleanup_thread.start()
        release.set()
        preview_thread.join(2); cleanup_thread.join(2)
        self.assertFalse(preview_thread.is_alive())
        self.assertFalse(cleanup_thread.is_alive())
        self.assertTrue(preview_result[0]["ok"])
        self.assertTrue(cleanup_result[0]["ok"])
        self.assertNotIn("T-1", restarted._task_sources)
        self.assertNotIn(token, restarted._selections)
        self.assertEqual(restarted.get_source_preview("T-1", "F-1")["code"], "SOURCE_RESELECT_REQUIRED")

    def test_human_save_merge_and_split_validate_before_atomic_write(self):
        self.seed_task()
        self.store.save_findings([finding("T-1", "F-2", fact_summary="虚构单位乙涉及金额 12 万元",
                                           source_excerpt="F-2 独立来源摘录", rationale="F-2 独立维度证据")])
        payload = asdict(finding("evil", "wrong", review_status="已接受"))
        saved = self.bridge.save_finding("T-1", "F-1", payload)
        self.assertEqual(saved["finding"]["task_id"], "T-1")
        self.assertEqual(saved["finding"]["finding_id"], "F-1")
        merged = self.bridge.merge_findings("T-1", ["F-1", "F-2"], asdict(finding("x", "x")))
        self.assertTrue(merged["ok"])
        stored = {item.finding_id: item for item in self.store.list_findings("T-1")}
        self.assertEqual(stored["F-2"].review_status, "已接受")
        self.assertEqual(stored["F-2"].merged_into, "F-1")
        self.assertEqual(stored["F-2"].fact_summary, "虚构单位乙涉及金额 12 万元")
        self.assertEqual(stored["F-2"].source_excerpt, "F-2 独立来源摘录")
        self.assertEqual(stored["F-2"].rationale, "F-2 独立维度证据")
        self.assertEqual(stored["F-1"].merged_finding_ids, ("F-2",))
        edit = asdict(stored["F-1"])
        edit.pop("merged_finding_ids"); edit.pop("merged_into")
        self.bridge.save_finding("T-1", "F-1", edit)
        self.assertEqual(self.store.list_findings("T-1")[0].merged_finding_ids, ("F-2",))
        before = [asdict(item) for item in self.store.list_findings("T-1")]
        invalid = self.bridge.split_finding("T-1", "F-1", [asdict(finding("x", "new")), {"finding_id": "bad"}])
        self.assertFalse(invalid["ok"])
        self.assertEqual([asdict(item) for item in self.store.list_findings("T-1")], before)
        split = self.bridge.split_finding("T-1", "F-1", [asdict(finding("x", "new-a")), asdict(finding("x", "new-b"))])
        self.assertTrue(split["ok"])
        by_id = {item.finding_id: item.review_status for item in self.store.list_findings("T-1")}
        self.assertEqual(by_id["F-1"], "已排除")
        self.assertEqual(by_id["new-a"], "待确认")

    def test_merge_then_exclude_commits_the_complete_lineage(self):
        from desktop.models import RiskDecision
        from desktop.workbook_writer import preview_changes, write_versioned_workbook

        self.seed_task(); self.store.save_findings([finding("T-1", "F-2")])
        payload = asdict(finding("T-1", "F-1", review_status="已接受"))
        self.assertTrue(self.bridge.merge_findings("T-1", ["F-1", "F-2"], payload)["ok"])
        payload["review_status"] = "已排除"
        self.assertTrue(self.bridge.save_finding("T-1", "F-1", payload)["ok"])
        stored = tuple(self.store.list_findings("T-1"))
        self.assertEqual({item.finding_id: item.review_status for item in stored}, {"F-1": "已排除", "F-2": "已排除"})
        decision = RiskDecision(action="exclude", finding_ids=("F-1", "F-2"))
        preview = preview_changes(self.workbook, (decision,), stored)
        result = write_versioned_workbook(self.workbook, (decision,), stored,
                                          expected_commit_token=preview["commit_token"],
                                          timestamp="20260904_0101", output_dir=self.root / "exclude-output")
        self.assertTrue(result.workbook_path.is_file())

    def test_merge_then_split_commits_children_without_orphaned_members(self):
        from desktop.models import RiskDecision
        from desktop.workbook_writer import preview_changes, write_versioned_workbook

        self.seed_task(); self.store.save_findings([finding("T-1", "F-2")])
        payload = asdict(finding("T-1", "F-1", review_status="已接受"))
        self.bridge.merge_findings("T-1", ["F-1", "F-2"], payload)
        child_a = asdict(finding("T-1", "F-1-A", matched_risk_id="", review_status="待确认"))
        child_b = asdict(finding("T-1", "F-1-B", matched_risk_id="", review_status="待确认"))
        self.assertTrue(self.bridge.split_finding("T-1", "F-1", [child_a, child_b])["ok"])
        for child_id in ("F-1-A", "F-1-B"):
            child = next(item for item in self.store.list_findings("T-1") if item.finding_id == child_id)
            self.bridge.save_finding("T-1", child_id, {**asdict(child), "review_status": "已接受"})
        stored = tuple(self.store.list_findings("T-1"))
        statuses = {item.finding_id: item.review_status for item in stored}
        self.assertEqual(statuses, {"F-1": "已排除", "F-1-A": "已接受", "F-1-B": "已接受", "F-2": "已排除"})
        common = dict(action="create", risk_id="", name="拆分风险", domain="采购与外包",
                      description="虚构拆分事实", owner_dept="采购部", period="2026H2", likelihood=3,
                      impact_scores={dim: 2 for dim in DIMS}, rationale="虚构拆分证据",
                      remediation_status="未整改")
        decisions = (RiskDecision(action="exclude", finding_ids=("F-1", "F-2")),
                     RiskDecision(finding_ids=("F-1-A",), **common),
                     RiskDecision(finding_ids=("F-1-B",), **common))
        preview = preview_changes(self.workbook, decisions, stored)
        result = write_versioned_workbook(self.workbook, decisions, stored,
                                          expected_commit_token=preview["commit_token"],
                                          timestamp="20260904_0102", output_dir=self.root / "split-output")
        self.assertTrue(result.workbook_path.is_file())

    def test_preview_and_commit_bind_selection_token_and_return_output_paths(self):
        self.seed_task()
        selected = {"selection_token": self.bind_workbook()}
        decision = self.create_decision()
        preview = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", [decision], "preview", True)
        self.assertEqual(preview["commit_token"], "token-1")
        bad = self.bridge.commit_to_workbook("T-1", selected["selection_token"], "2026H1", [decision], "wrong")
        self.assertEqual(bad["code"], "PREVIEW_REQUIRED")
        period_risk = {"risk_id": "R001", "name": "合成风险", "domain": "采购与外包", "description": "合成事实", "owner_dept": "审计部", "period": "2026H1", "likelihood": 3, **{dim: 2 for dim in DIMS}, "rationale": "合成依据"}
        with patch("desktop.bridge.load_dataset", return_value=({}, [period_risk], [])):
            done = self.bridge.commit_to_workbook("T-1", selected["selection_token"], "2026H1", [decision], "token-1")
        self.assertEqual(done["workbook_path"], str(Path("C:/synthetic/out.xlsx")))
        self.assertEqual(done["export_dir"], str(Path("C:/synthetic/export")))
        self.assertEqual(done["period_data"], {"period": "2026H1", "risks": [period_risk], "controls": []})

    def test_cleanup_only_forgets_runtime_mappings(self):
        self.seed_task()
        selected = self.select(self.report)
        self.bridge._task_sources["T-1"] = (self.report, selected["selection_token"])
        result = self.bridge.cleanup_task("T-1")
        self.assertTrue(result["ok"])
        self.assertEqual(self.store.get_task("T-1").task_id, "T-1")
        self.assertTrue(self.report.exists())
        self.assertNotIn("T-1", self.bridge._task_sources)

    def test_cleanup_task_is_scoped_to_its_report_and_workbook_preview(self):
        second_report = self.root / "second.pdf"; second_report.write_bytes(b"second synthetic report")
        self.seed_task("T-1", self.report)
        self.seed_task("T-2", second_report)
        report_one = self.select(self.report)["selection_token"]
        report_two = self.select(second_report)["selection_token"]
        self.bridge._task_sources.update({"T-1": (self.report, report_one), "T-2": (second_report, report_two)})
        first_book = self.bind_workbook("T-1")
        second_book_file = self.root / "second.xlsx"; shutil.copy2(self.workbook, second_book_file)
        second_book = self.bind_workbook("T-2", path=second_book_file)
        decision = [self.create_decision()]
        self.bridge.preview_commit("T-1", first_book, "2026H1", decision, "preview", True)
        self.bridge.preview_commit("T-2", second_book, "2026H1", decision, "preview", True)
        cleaned = self.bridge.cleanup_task("T-1")
        self.assertTrue(cleaned["ok"])
        self.assertEqual(self.pipeline.cleaned, ["T-1"])
        self.assertNotIn("T-1", self.bridge._task_sources)
        self.assertNotIn(report_one, self.bridge._selections)
        self.assertNotIn(first_book, self.bridge._selections)
        self.assertNotIn(first_book, self.bridge._commit_previews)
        self.assertEqual(self.bridge._task_sources["T-2"], (second_report, report_two))
        self.assertIn(report_two, self.bridge._selections)
        self.assertIn(second_book, self.bridge._selections)
        self.assertIn(second_book, self.bridge._commit_previews)

    def test_real_bridge_response_keys_match_desktop_script_contract(self):
        self.seed_task()
        self.store.save_findings([finding("T-1", "F-2")])
        merged = self.bridge.merge_findings("T-1", ["F-1", "F-2"], asdict(finding("T-1", "F-1", review_status="已接受")))
        self.assertEqual(set(merged) - {"ok"}, {"findings"})
        self.assertEqual({item["finding_id"] for item in merged["findings"]}, {"F-1", "F-2"})
        split = self.bridge.split_finding("T-1", "F-1", [asdict(finding("T-1", "F-1-A")), asdict(finding("T-1", "F-1-B"))])
        self.assertEqual(set(split) - {"ok"}, {"findings"})
        self.assertEqual({item["finding_id"] for item in split["findings"]}, {"F-1", "F-2", "F-1-A", "F-1-B"})
        selected = {"selection_token": self.bind_workbook()}
        decision = [self.create_decision("F-1-A")]
        preview = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", decision, "preview", True)
        self.assertEqual(set(preview) - {"ok"}, {"commit_token", "new_risks", "updated_risks", "new_controls", "excluded_count", "warnings"})
        risk = {"risk_id": "R001", "name": "合成风险", "domain": "采购与外包", "description": "合成事实", "owner_dept": "审计部", "period": "2026H1", "likelihood": 3, **{dim: 2 for dim in DIMS}, "rationale": "合成依据"}
        with patch("desktop.bridge.load_dataset", return_value=({}, [risk], [])):
            committed = self.bridge.commit_to_workbook("T-1", selected["selection_token"], "2026H1", decision, preview["commit_token"])
            self.assertEqual(set(committed) - {"ok"}, {"workbook_path", "export_dir", "period_data"})

    def configure_catalog(self, entity_name="虚构主体甲"):
        catalog_root = self.root / "catalog"
        catalog_root.mkdir(exist_ok=True)
        self.window.selected = [str(catalog_root)]
        selected = self.bridge.choose_catalog_root()
        self.assertTrue(selected["ok"])
        configured = self.bridge.configure_workspace(selected["selection_token"], entity_name)
        self.assertTrue(configured["ok"])
        return catalog_root, configured

    def test_catalog_workspace_is_selected_once_and_returned_by_bootstrap(self):
        catalog_root, configured = self.configure_catalog()
        self.assertEqual(configured["workspace"]["entity_name"], "虚构主体甲")
        self.assertEqual(Path(self.store.get_setting("catalog_root")), catalog_root.resolve())
        bootstrap = self.bridge.get_bootstrap()
        self.assertEqual(bootstrap["workspace"]["entity_name"], "虚构主体甲")
        self.assertEqual(bootstrap["catalog_reports"], [])
        self.assertEqual(self.window.calls[-1]["allow_multiple"], False)

    def test_reviewed_task_is_atomically_saved_to_catalog_then_removed_from_sqlite(self):
        self.configure_catalog()
        self.seed_task("T-catalog")
        accepted = finding("T-catalog", review_status="已接受")
        self.store.save_findings([accepted])
        saved = self.bridge.save_report_to_catalog("T-catalog", {
            "audit_project": "采购专项审计",
            "report_title": "采购管理专项审计报告",
            "report_date": "2026-08-28",
        })
        self.assertTrue(saved["ok"])
        self.assertEqual(saved["report"]["audit_project"], "采购专项审计")
        self.assertIsNone(self.store.get_task("T-catalog"))
        self.assertEqual(self.store.list_findings("T-catalog"), [])
        catalog = self.bridge.list_catalog_reports()
        self.assertEqual(catalog["reports"][0]["finding_count"], 1)
        self.assertNotIn(str(self.report), str(catalog))

    def test_catalog_rejects_pending_review_and_supports_trash_and_clear(self):
        self.configure_catalog()
        self.seed_task("T-pending")
        pending = self.bridge.save_report_to_catalog("T-pending", {
            "audit_project": "采购审计", "report_title": "待复核报告", "report_date": "",
        })
        self.assertEqual(pending["code"], "REPORT_REVIEW_INCOMPLETE")
        for task_id, title in (("T-one", "报告一"), ("T-two", "报告二")):
            self.seed_task(task_id)
            self.store.save_findings([finding(task_id, review_status="已接受")])
            self.assertTrue(self.bridge.save_report_to_catalog(task_id, {
                "audit_project": "采购审计", "report_title": title, "report_date": "",
            })["ok"])
        reports = self.bridge.list_catalog_reports()["reports"]
        trashed = self.bridge.trash_catalog_report(reports[0]["report_id"])
        self.assertTrue(trashed["ok"])
        cleared = self.bridge.clear_catalog_reports()
        self.assertEqual(cleared["count"], 1)
        self.assertEqual(self.bridge.list_catalog_reports()["reports"], [])

    def test_catalog_batch_clones_exact_report_versions_and_preserves_source_labels(self):
        catalog_root, _ = self.configure_catalog()
        from desktop.catalog import CatalogStore
        catalog = CatalogStore(catalog_root)
        first = catalog.save_report(
            AnalysisTask("T-A", "a.pdf", "a" * 64, "2026-09-01T00:00:00Z", "待复核", "synthetic", "text"),
            [finding("T-A", "F-A", review_status="已接受")],
            audit_project="采购专项审计", report_title="报告甲",
        )
        invalid = finding("T-B", "F-B", review_status="已接受")
        invalid.matched_risk_id = "R999"
        second = catalog.save_report(
            AnalysisTask("T-B", "b.pdf", "b" * 64, "2026-09-02T00:00:00Z", "待复核", "synthetic", "ocr"),
            [invalid], audit_project="供应链内控审计", report_title="报告乙",
        )
        workbook_token = self.select(self.workbook, "workbook")["selection_token"]
        batch = self.bridge.create_catalog_batch(
            [first["report_id"], second["report_id"]], workbook_token, "2026H2"
        )
        self.assertTrue(batch["ok"], batch)
        self.assertEqual(len(batch["reports"]), 2)
        cloned = batch["findings"]
        self.assertEqual(len(cloned), 2)
        remap = next(item for item in cloned if "报告乙" in item["source_page"])
        self.assertEqual(remap["matched_risk_id"], "")
        self.assertEqual(remap["review_status"], "待确认")
        source = self.bridge.get_source_preview(batch["task"]["task_id"], remap["finding_id"])
        self.assertEqual(source["kind"], "text")
        self.assertEqual(source["source_report_title"], "报告乙")
        refs = batch["report_refs"]
        self.assertEqual({item["recognition_version"] for item in refs}, {1})

    def test_completed_catalog_batch_writes_immutable_batch_snapshot(self):
        catalog_root, _ = self.configure_catalog()
        from desktop.catalog import CatalogStore
        catalog = CatalogStore(catalog_root)
        report = catalog.save_report(
            AnalysisTask("T-A", "a.pdf", "a" * 64, "2026-09-01T00:00:00Z", "待复核", "synthetic", "text"),
            [finding("T-A", "F-A", review_status="已接受")],
            audit_project="采购专项审计", report_title="报告甲",
        )
        workbook_token = self.select(self.workbook, "workbook")["selection_token"]
        batch = self.bridge.create_catalog_batch([report["report_id"]], workbook_token, "2026H2")
        self.assertTrue(batch["ok"], batch)
        finding_id = batch["findings"][0]["finding_id"]
        decision = self.create_decision(finding_id, period="2026H2")
        preview = self.bridge.preview_commit(batch["task"]["task_id"], workbook_token, "2026H2", [decision], "preview", True)
        self.assertTrue(preview["ok"])
        period_risk = {"risk_id": "R001", "name": "虚构风险", "domain": "采购与外包", "description": "虚构事实", "owner_dept": "审计部", "period": "2026H2", "likelihood": 3, **{dim: 2 for dim in DIMS}, "rationale": "虚构依据"}
        with patch("desktop.bridge.load_dataset", return_value=({}, [period_risk], [])):
            committed = self.bridge.commit_to_workbook(batch["task"]["task_id"], workbook_token, "2026H2", [decision], preview["commit_token"])
        self.assertTrue(committed["ok"])
        snapshots = list((catalog_root / "batches").glob("*/batch.json"))
        self.assertEqual(len(snapshots), 1)
        snapshot = CatalogStore(catalog_root).load_batch(snapshots[0].parent.name)
        self.assertEqual(snapshot["report_refs"][0]["report_id"], report["report_id"])
        self.assertEqual(snapshot["period"], "2026H2")

    def test_preview_rejects_decisions_outside_the_selected_period(self):
        self.seed_task()
        selected = {"selection_token": self.bind_workbook()}
        wrong_period = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", [self.create_decision(period="2026H2")])
        self.assertFalse(wrong_period["ok"])
        self.assertEqual(self.writer.preview_calls, [])

    def test_preview_loads_current_controls_before_confirmed_final_preview(self):
        self.seed_task()
        selected = {"selection_token": self.bind_workbook()}
        identities = [{"finding_ids": ["F-1"], "action": "merge", "risk_id": "R001", "period": "2026H1"}]
        current = [{"description": "保留的合成控制", "score": 4, "key": True}]
        with patch("desktop.bridge.load_current_controls", return_value=[{**identities[0], "controls": current}]) as loaded:
            stage = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", identities, "load_controls")
        self.assertEqual(stage, {"ok": True, "controls_by_decision": [{**identities[0], "controls": current}]})
        loaded.assert_called_once()
        denied = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", [self.create_decision()], "preview", False)
        self.assertFalse(denied["ok"])
        final = self.bridge.preview_commit("T-1", selected["selection_token"], "2026H1", [self.create_decision()], "preview", True)
        self.assertTrue(final["ok"])
        self.assertEqual(final["commit_token"], "token-1")

    def test_real_save_accept_merge_and_split_use_nested_impact_scores_only(self):
        self.seed_task()
        self.store.save_findings([finding("T-1", "F-2")])
        payload = asdict(finding("ignored", "ignored", review_status="已接受"))
        self.assertIn("impact_scores", payload)
        self.assertFalse(any(dim in payload for dim in DIMS))
        saved = self.bridge.save_finding("T-1", "F-1", payload)
        self.assertEqual(saved["finding"]["review_status"], "已接受")
        merged = self.bridge.merge_findings("T-1", ["F-1", "F-2"], payload)
        self.assertEqual({item["finding_id"] for item in merged["findings"]}, {"F-1", "F-2"})
        split = self.bridge.split_finding("T-1", "F-1", [asdict(finding("ignored", "F-1-A")), asdict(finding("ignored", "F-1-B"))])
        self.assertEqual({item["finding_id"] for item in split["findings"]}, {"F-1", "F-2", "F-1-A", "F-1-B"})
        extra = {**payload, "imp_financial": 2}
        rejected = self.bridge.save_finding("T-1", "F-1-A", extra)
        self.assertFalse(rejected["ok"])
        self.assertEqual(rejected["code"], "VALIDATION_ERROR")


class AppBootstrapTests(unittest.TestCase):
    def test_build_bridge_does_not_load_a_packaged_example_risk_catalog(self):
        from desktop import app

        class Credentials(CredentialMemory):
            def assert_windows_backend(self): pass

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            bridge = app.build_bridge(
                state_path=root / "state.db", temp_path=root / "tasks",
                credential_store=Credentials(),
                resource_provider=lambda name: (_ for _ in ()).throw(AssertionError(f"packaged resource read: {name}")),
            )
            try:
                self.assertEqual(bridge.get_bootstrap()["risk_catalog"], [])
            finally:
                bridge._pipeline.close()

    def test_main_creates_private_edge_window_and_closes_pipeline(self):
        from desktop import app

        class PipelineClose:
            def __init__(self): self.closed = False
            def close(self): self.closed = True

        class Bridge:
            def __init__(self): self.pipeline = PipelineClose(); self.window = None; self.webview = None
            def __getattr__(self, name):
                if name == "attach_window":
                    def attach(window, webview_module=None):
                        self.window, self.webview = window, webview_module
                    return attach
                raise AttributeError(name)

        class Event:
            def __init__(self): self.callbacks = []
            def __iadd__(self, callback): self.callbacks.append(callback); return self

        class Window:
            def __init__(self): self.events = type("Events", (), {"closed": Event()})()

        class Webview:
            def __init__(self): self.calls = []; self.window = Window()
            def create_window(self, **kwargs): self.calls.append(kwargs); return self.window
            def start(self, **kwargs):
                self.calls.append(("start", kwargs))
                for callback in self.window.events.closed.callbacks: callback()

        webview, bridge = Webview(), Bridge()
        app.main(webview_module=webview, bridge_factory=lambda: bridge,
                 resource_provider=lambda _: Path("C:/synthetic/risk_heatmap.html"))
        created = webview.calls[0]
        self.assertEqual(created["title"], "审计风险评估热力图谱")
        self.assertEqual(created["url"], Path("C:/synthetic/risk_heatmap.html").as_uri())
        self.assertEqual((created["width"], created["height"], created["min_size"]), (1440, 920, (1120, 720)))
        self.assertIs(created["js_api"], bridge)
        self.assertIs(bridge.window, webview.window)
        self.assertIs(bridge.webview, webview)
        self.assertEqual(webview.calls[1], ("start", {"gui": "edgechromium", "private_mode": True, "debug": False}))
        self.assertTrue(bridge.pipeline.closed)


if __name__ == "__main__":
    unittest.main()
