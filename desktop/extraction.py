"""Local PDF and DOCX extraction with explicit OCR and vision-review routing."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import math
from pathlib import Path
from typing import Protocol
from pathlib import PurePosixPath
import re
import zipfile

import pypdfium2
from PIL import Image
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from desktop.models import ExtractedBlock
from desktop.ocr import OcrError


# Resource caps prevent malformed files from exhausting local desktop resources.
MAX_SOURCE_BYTES = 50 * 1024 * 1024
MAX_PDF_PAGES = 200
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200
MAX_IMAGE_BYTES = 20 * 1024 * 1024
MAX_IMAGE_PIXELS = 40_000_000
MAX_RENDER_PIXELS = 25_000_000
PDF_RENDER_SCALE = 2.0
MAX_DOCX_PREVIEW_CHARS = 20_000
_DOCX_PREVIEW_LOCATOR = re.compile(r"^Word (段落|表格) ([1-9][0-9]*)$")


class OcrReader(Protocol):
    def read(self, image_path: Path) -> str: ...


class ExtractionError(RuntimeError):
    def __init__(self, code: str, message: str) -> None:
        self.code = code
        self.message = message
        super().__init__(message)


@dataclass(slots=True)
class ExtractionResult:
    blocks: list[ExtractedBlock]
    method: str


def text_is_usable(text: str) -> bool:
    compact = "".join(text.split()) if isinstance(text, str) else ""
    if len(compact) < 40:
        return False
    printable = sum(char.isprintable() and char != "\ufffd" for char in compact)
    return printable / len(compact) >= 0.90


def _safe_temp_root(temp_dir: Path) -> Path:
    root = Path(temp_dir)
    if not root.exists() or not root.is_dir():
        raise ExtractionError("TEMP_DIR_INVALID", "任务临时目录不可用")
    return root.resolve()


def _source_size_ok(source: Path) -> None:
    if source.stat().st_size > MAX_SOURCE_BYTES:
        raise ExtractionError("SOURCE_TOO_LARGE", "输入文件超过安全大小限制")


def _output_path(root: Path, filename: str) -> Path:
    target = (root / filename).resolve()
    if root not in target.parents:
        raise ExtractionError("TEMP_DIR_INVALID", "任务临时目录不可用")
    return target


def _close(resource: object | None) -> None:
    close = getattr(resource, "close", None)
    if callable(close):
        close()


def _validate_image_bytes(data: bytes) -> None:
    if len(data) > MAX_IMAGE_BYTES:
        raise ExtractionError("IMAGE_BYTES_LIMIT", "图片超过安全大小限制")
    try:
        with Image.open(BytesIO(data)) as image:
            width, height = image.size
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("IMAGE_READ_FAILED", "无法读取嵌入图片") from exc
    if width * height > MAX_IMAGE_PIXELS:
        raise ExtractionError("IMAGE_PIXELS_LIMIT", "图片像素超过安全限制")


def _validate_image_file(image_path: Path) -> None:
    try:
        _validate_image_bytes(image_path.read_bytes())
    except ExtractionError:
        raise
    except OSError as exc:
        raise ExtractionError("IMAGE_READ_FAILED", "无法读取渲染图片") from exc


def _read_ocr(ocr: OcrReader, image_path: Path, locator: str) -> ExtractedBlock:
    _validate_image_file(image_path)
    try:
        text = ocr.read(image_path)
    except OcrError:
        return ExtractedBlock(locator, "需要视觉复核：本地OCR不可用或识别失败。", "vision_required", True, str(image_path))
    if text_is_usable(text):
        return ExtractedBlock(locator, text, "ocr", image_path=str(image_path))
    return ExtractedBlock(locator, "需要视觉复核：OCR文本质量不足。", "vision_required", True, str(image_path))


def _render_pdf_page(page: object, output_path: Path) -> None:
    bitmap = None
    image = None
    try:
        bitmap = page.render(scale=PDF_RENDER_SCALE)
        image = bitmap.to_pil()
        image.save(output_path, format="PNG")
    finally:
        if image is not None:
            image.close()
        _close(bitmap)


def _check_pdf_render_size(page: object) -> None:
    try:
        width, height = page.get_size()
    except Exception as exc:
        raise ExtractionError("PDF_READ_FAILED", "无法读取PDF页面尺寸") from exc
    pixels = math.ceil(width * PDF_RENDER_SCALE) * math.ceil(height * PDF_RENDER_SCALE)
    if pixels > MAX_RENDER_PIXELS:
        raise ExtractionError("PDF_RENDER_LIMIT", "PDF页面渲染像素超过安全限制")


def _extract_pdf(path: Path, temp_root: Path, ocr: OcrReader) -> list[ExtractedBlock]:
    document = None
    try:
        document = pypdfium2.PdfDocument(path)
        if len(document) > MAX_PDF_PAGES:
            raise ExtractionError("PDF_PAGE_LIMIT", "PDF页数超过安全限制")
        blocks: list[ExtractedBlock] = []
        for index in range(len(document)):
            page = text_page = None
            locator = f"第 {index + 1} 页"
            try:
                page = document[index]
                try:
                    text_page = page.get_textpage()
                    text = text_page.get_text_bounded()
                except Exception:
                    text = ""
                if text_is_usable(text):
                    blocks.append(ExtractedBlock(locator, text, "text"))
                else:
                    _check_pdf_render_size(page)
                    image_path = _output_path(temp_root, f"pdf_page_{index + 1:04d}.png")
                    _render_pdf_page(page, image_path)
                    blocks.append(_read_ocr(ocr, image_path, locator))
            finally:
                _close(text_page)
                _close(page)
        return blocks
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("PDF_READ_FAILED", f"无法读取PDF文件：{path.name}") from exc
    finally:
        _close(document)


def _image_suffix(content_type: str) -> str:
    return {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
    }.get(content_type, ".img")


def _zip_name(name: str) -> str:
    normalized = name.replace("\\", "/")
    path = PurePosixPath(normalized)
    if path.is_absolute() or any(part == ".." for part in path.parts) or (path.parts and ":" in path.parts[0]):
        raise ExtractionError("DOCX_ENTRY_PATH_INVALID", "Word压缩包包含不安全条目")
    return "/".join(part for part in path.parts if part != ".")


def _preflight_docx(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
    except (OSError, zipfile.BadZipFile) as exc:
        raise ExtractionError("DOCX_READ_FAILED", f"无法读取Word文件：{path.name}") from exc
    if len(entries) > MAX_DOCX_ENTRIES:
        raise ExtractionError("DOCX_ENTRY_LIMIT", "Word压缩包条目超过安全限制")
    total = 0
    names: set[str] = set()
    for entry in entries:
        canonical = _zip_name(entry.filename)
        if not canonical or canonical in names:
            raise ExtractionError("DOCX_DUPLICATE_ENTRY", "Word压缩包包含重复条目")
        names.add(canonical)
        total += entry.file_size
        if entry.file_size and (entry.compress_size == 0 or entry.file_size / entry.compress_size > MAX_DOCX_COMPRESSION_RATIO):
            raise ExtractionError("DOCX_COMPRESSION_RATIO", "Word压缩包压缩比超过安全限制")
    if total > MAX_DOCX_UNCOMPRESSED_BYTES:
        raise ExtractionError("DOCX_SIZE_LIMIT", "Word压缩包解压大小超过安全限制")


def _image_paths(drawing: object, part: object, temp_root: Path, image_sequence: list[int]) -> list[Path]:
    images: list[Path] = []
    for blip in drawing.xpath(".//a:blip"):
        relationship_id = blip.get(qn("r:embed"))
        if not relationship_id:
            continue
        relationship = part.rels.get(relationship_id)
        if relationship is None or not hasattr(relationship.target_part, "blob"):
            continue
        suffix = _image_suffix(getattr(relationship.target_part, "content_type", ""))
        data = relationship.target_part.blob
        _validate_image_bytes(data)
        image_sequence[0] += 1
        image_path = _output_path(temp_root, f"word_image_{image_sequence[0]:04d}{suffix}")
        image_path.write_bytes(data)
        images.append(image_path)
    return images


def _inline_events(element: object, part: object, temp_root: Path, image_sequence: list[int]) -> list[str | Path]:
    events: list[str | Path] = []
    tag = getattr(element, "tag", "")
    if tag.endswith("}drawing"):
        return _image_paths(element, part, temp_root, image_sequence)
    if tag.endswith("}t"):
        return [getattr(element, "text", "") or ""]
    if tag.endswith("}tab"):
        return ["\t"]
    if tag.endswith("}br") or tag.endswith("}cr"):
        return ["\n"]
    for child in element.iterchildren():
        events.extend(_inline_events(child, part, temp_root, image_sequence))
    return events


def _paragraph_events(paragraph: Paragraph, temp_root: Path, image_sequence: list[int]) -> list[str | Path]:
    events: list[str | Path] = []
    for child in paragraph._p.iterchildren():
        events.extend(_inline_events(child, paragraph.part, temp_root, image_sequence))
    return events


def _append_events(blocks: list[ExtractedBlock], locator: str, events: list[str | Path], ocr: OcrReader) -> None:
    text = ""
    for event in events:
        if isinstance(event, str):
            text += event
            continue
        if text.strip():
            blocks.append(ExtractedBlock(locator, text.strip(), "text"))
        text = ""
        blocks.append(_read_ocr(ocr, event, locator))
    if text.strip():
        blocks.append(ExtractedBlock(locator, text.strip(), "text"))


def _table_events(table: Table, temp_root: Path, image_sequence: list[int]) -> list[str | Path]:
    events: list[str | Path] = []
    rows = table._tbl.tr_lst
    for row_index, row in enumerate(rows):
        cells = row.tc_lst
        for cell_index, cell_xml in enumerate(cells):
            cell = _Cell(cell_xml, table)
            child_groups: list[list[str | Path]] = []
            for child in cell._tc.iterchildren():
                if child.tag.endswith("}p"):
                    group = _paragraph_events(Paragraph(child, cell), temp_root, image_sequence)
                elif child.tag.endswith("}tbl"):
                    group = _table_events(Table(child, cell), temp_root, image_sequence)
                else:
                    continue
                if group:
                    child_groups.append(group)
            for group_index, group in enumerate(child_groups):
                if group_index:
                    events.append("\n")
                events.extend(group)
            if cell_index < len(cells) - 1:
                events.append("\t")
        if row_index < len(rows) - 1:
            events.append("\n")
    return events


def _extract_docx(path: Path, temp_root: Path, ocr: OcrReader) -> list[ExtractedBlock]:
    try:
        _preflight_docx(path)
        document = Document(path)
        blocks: list[ExtractedBlock] = []
        paragraph_count = table_count = 0
        image_sequence = [0]
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph_count += 1
                paragraph = Paragraph(child, document)
                _append_events(blocks, f"Word 段落 {paragraph_count}", _paragraph_events(paragraph, temp_root, image_sequence), ocr)
            elif child.tag.endswith("}tbl"):
                table_count += 1
                _append_events(blocks, f"Word 表格 {table_count}", _table_events(Table(child, document), temp_root, image_sequence), ocr)
        return blocks
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("DOCX_READ_FAILED", f"无法读取Word文件：{path.name}") from exc


def extract_docx_source_text(path: Path, locator: str) -> str:
    """Re-read one structural DOCX location from a hash-checked snapshot."""
    match = _DOCX_PREVIEW_LOCATOR.fullmatch(locator) if isinstance(locator, str) else None
    if match is None:
        raise ExtractionError("DOCX_LOCATOR_INVALID", "Word来源位置格式无效")
    source = Path(path)
    _source_size_ok(source)
    _preflight_docx(source)
    kind, wanted = match.group(1), int(match.group(2))
    try:
        document = Document(source)
        count = 0
        for child in document.element.body.iterchildren():
            if kind == "段落" and child.tag.endswith("}p"):
                count += 1
                if count == wanted:
                    text = Paragraph(child, document).text.strip()
                    break
            elif kind == "表格" and child.tag.endswith("}tbl"):
                count += 1
                if count == wanted:
                    table = Table(child, document)
                    text = "\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows).strip()
                    break
        else:
            raise ExtractionError("DOCX_LOCATOR_NOT_FOUND", "Word来源位置不存在")
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError("DOCX_READ_FAILED", f"无法读取Word文件：{source.name}") from exc
    if not text:
        raise ExtractionError("DOCX_PREVIEW_EMPTY", "Word来源位置没有可显示文本")
    return text[:MAX_DOCX_PREVIEW_CHARS]


def _aggregate_method(blocks: list[ExtractedBlock]) -> str:
    methods = {block.method for block in blocks}
    if len(methods) == 1:
        return next(iter(methods), "text")
    return "mixed"


def extract_report(path: Path, temp_dir: Path, ocr: OcrReader) -> ExtractionResult:
    source = Path(path)
    if not source.exists():
        raise ExtractionError("INPUT_NOT_FOUND", f"输入文件不存在：{source.name}")
    if not source.is_file():
        raise ExtractionError("INPUT_NOT_FILE", f"输入不是文件：{source.name}")
    _source_size_ok(source)
    temp_root = _safe_temp_root(temp_dir)
    suffix = source.suffix.lower()
    if suffix == ".doc":
        raise ExtractionError("DOC_LEGACY_UNSUPPORTED", f"不支持旧版Word文件：{source.name}")
    if suffix == ".pdf":
        blocks = _extract_pdf(source, temp_root, ocr)
    elif suffix == ".docx":
        blocks = _extract_docx(source, temp_root, ocr)
    else:
        raise ExtractionError("FILE_TYPE_UNSUPPORTED", f"不支持的文件类型：{source.name}")
    return ExtractionResult(blocks, _aggregate_method(blocks))
